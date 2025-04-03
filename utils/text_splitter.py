# utils/text_splitter.py

import json
import os
import re
from typing import List, Dict, Any, Optional

import yaml
from tqdm import tqdm
from docx import Document

class DocumentSplitter:
    """
    文档分割器，将文档切分为小块
    """
    
    def __init__(self, config_path: str = "../config.yaml"):
        """
        初始化文档分割器
        
        Args:
            config_path: 配置文件路径
        """
        # 加载配置
        with open(config_path, 'r', encoding='utf-8') as f:
            self.config = yaml.safe_load(f)
        
        # 分割参数 - 使用默认值，如果配置中没有相应的项
        if 'splitter' not in self.config:
            print("警告: 配置文件中缺少 'splitter' 配置项，使用默认值")
            self.chunk_size = 1000
            self.chunk_overlap = 200
        else:
            self.chunk_size = self.config['splitter'].get('chunk_size', 1000)
            self.chunk_overlap = self.config['splitter'].get('chunk_overlap', 200)
    
    def process_directory(self, directory: str, output_file: str) -> List[Dict[str, Any]]:
        """
        处理目录中的所有文档
        
        Args:
            directory: 文档目录
            output_file: 输出文件路径
            
        Returns:
            处理后的文档列表
        """
        documents = []
        
        # 确保目录存在
        if not os.path.exists(directory):
            print(f"目录不存在: {directory}")
            os.makedirs(directory, exist_ok=True)
            return documents
        
        # 获取所有文件
        files = []
        for root, _, filenames in os.walk(directory):
            for filename in filenames:
                files.append(os.path.join(root, filename))
        
        # 处理文件
        for file_path in tqdm(files, desc="Processing files"):
            try:
                # 根据文件扩展名处理
                ext = os.path.splitext(file_path)[1].lower()
                
                if ext == '.txt' or ext == '.md':
                    # 处理文本文件
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                    chunks = self.split_text(text)
                    
                elif ext == '.docx':
                    # 处理 Word 文档
                    text = self.read_docx(file_path)
                    chunks = self.split_text(text)
                    
                else:
                    # 跳过不支持的文件类型
                    print(f"跳过不支持的文件类型: {file_path}")
                    continue
                
                # 添加到文档列表
                for i, chunk in enumerate(chunks):
                    doc_id = f"{os.path.basename(file_path)}_{i}"
                    documents.append({
                        "id": doc_id,
                        "content": chunk,
                        "metadata": {
                            "source": file_path,
                            "chunk_id": i
                        }
                    })
                    
            except Exception as e:
                print(f"处理文件时出错 {file_path}: {e}")
        
        # 保存结果
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(documents, f, ensure_ascii=False, indent=2)
        
        return documents
    
    def read_docx(self, file_path: str) -> str:
        """
        读取 Word 文档内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档内容
        """
        try:
            doc = Document(file_path)  # 使用 Document 而不是 docx.Document
            full_text = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():  # 跳过空段落
                    full_text.append(para.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():  # 跳过空单元格
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            print(f"读取 Word 文档时出错 {file_path}: {e}")
            return ""
    
    def split_text(self, text: str) -> List[str]:
        """
        将文本分割为小块
        
        Args:
            text: 输入文本
            
        Returns:
            文本块列表
        """
        # 如果文本为空，返回空列表
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        
        while start < len(text):
            # 计算当前块的结束位置
            end = min(start + self.chunk_size, len(text))
            
            # 如果不是最后一块，尝试在句子边界处分割
            if end < len(text):
                # 尝试在句号、问号、感叹号后分割
                for punct in ['. ', '? ', '! ', '\n\n']:
                    last_punct = text.rfind(punct, start, end)
                    if last_punct != -1:
                        end = last_punct + 1  # 包含标点符号
                        break
            
            # 添加当前块
            chunk = text[start:end].strip()
            if chunk:  # 跳过空块
                chunks.append(chunk)
            
            # 更新起始位置，考虑重叠
            start = end - self.chunk_overlap
            
            # 确保起始位置前进
            if start <= 0 or start >= len(text) - 1:
                break
        
        return chunks